from pathlib import Path

import pdfplumber
from docx import Document


def extract_resume_text(file_path: str) -> str:
    path = Path(file_path)

    if path.suffix.lower() == ".pdf":
        return extract_pdf_text(path)

    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if path.suffix.lower() == ".docx":
        return extract_docx_text(path)

    raise ValueError("Unsupported resume format")


def extract_pdf_text(path: Path) -> str:
    pages: list[str] = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text)

    return "\n\n".join(pages).strip()


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
